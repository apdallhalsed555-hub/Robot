import re
import os
import base64
import urllib.request
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MD_PATH = r"x:\Robot-main\Robot-main\Mahmoud_Contribution.md"
OUT_PATH = r"x:\Robot-main\Robot-main\Mahmoud_Contribution.docx"

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_code_block(doc, code_text, language=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F2F2")

    tc_pr = cell._tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '  <w:top w:w="60" w:type="dxa"/>'
        '  <w:left w:w="120" w:type="dxa"/>'
        '  <w:bottom w:w="60" w:type="dxa"/>'
        '  <w:right w:w="120" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tc_pr.append(margins)

    cell.paragraphs[0].clear()
    for i, line in enumerate(code_text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.style = doc.styles["Normal"]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(13)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph()

def add_mermaid_image(doc, mermaid_code, img_index):
    # Encode for mermaid.ink
    b64 = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
    url = f"https://mermaid.ink/img/{b64}"
    img_path = f"x:\\Robot-main\\Robot-main\\scratch\\mermaid_{img_index}.png"
    
    print(f"Downloading mermaid diagram {img_index}...")
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    try:
        with urllib.request.urlopen(req) as response, open(img_path, 'wb') as out_file:
            out_file.write(response.read())
        
        # Add to document
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(img_path, width=Inches(6.0))
        
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_run = caption.add_run(f"Figure: Architecture Diagram")
        c_run.font.italic = True
        c_run.font.size = Pt(9)
        c_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()
    except Exception as e:
        print(f"Failed to download mermaid {img_index}: {e}")
        add_code_block(doc, mermaid_code, "mermaid")

def add_formatted_paragraph(doc, text, style="Normal", bold_whole=False):
    p = doc.add_paragraph(style=style)
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`|[^*`]+)')
    parts = pattern.findall(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            run = p.add_run(part)
        if bold_whole:
            run.bold = True
    return p

def add_table_from_md(doc, header_line, rows_lines):
    def parse_row(line):
        return [c.strip() for c in line.strip().strip("|").split("|")]

    headers = parse_row(header_line)
    rows = [parse_row(r) for r in rows_lines]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(re.sub(r'\*\*(.+?)\*\*', r'\1', h))
        run.bold = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2C3E50")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci >= len(headers): continue
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', val)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            clean = re.sub(r'`(.+?)`', r'\1', clean)
            run = p.add_run(clean)
            run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F8F9FA")
    doc.add_paragraph()

def add_important_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF3CD")

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:left w:val="single" w:sz="12" w:space="0" w:color="FFC107"/>'
        '</w:tcBorders>'
    )
    tc_pr.append(borders)
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '  <w:top w:w="60" w:type="dxa"/>'
        '  <w:left w:w="120" w:type="dxa"/>'
        '  <w:bottom w:w="60" w:type="dxa"/>'
        '  <w:right w:w="120" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tc_pr.append(margins)

    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    run = p.add_run("⚠ IMPORTANT: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x85, 0x6D, 0x0D)
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    clean = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', clean)
    run2 = p.add_run(clean)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x85, 0x6D, 0x0D)
    doc.add_paragraph()

def create_toc_field(p):
    """Inserts a table of contents field in the paragraph."""
    run = p.add_run()
    fldChar1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    instrText = parse_xml(r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">TOC \o "1-3" \h \z \u</w:instrText>')
    fldChar2 = parse_xml(r'<w:fldChar w:fldCharType="separate" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    fldChar3 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def convert_md_to_docx():
    with open(MD_PATH, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    i = 0
    in_code_block = False
    code_buffer = []
    code_lang = ""
    in_table = False
    table_header = ""
    table_rows = []
    img_counter = 0

    while i < len(lines):
        line = lines[i].rstrip("\n")

        if line.startswith("```") and not in_code_block:
            if in_table:
                add_table_from_md(doc, table_header, table_rows)
                in_table = False; table_rows = []
            code_lang = line[3:].strip()
            in_code_block = True
            code_buffer = []
            i += 1
            continue

        if line.startswith("```") and in_code_block:
            code_text = "\n".join(code_buffer)
            if code_lang == "mermaid":
                img_counter += 1
                add_mermaid_image(doc, code_text, img_counter)
            else:
                add_code_block(doc, code_text, code_lang)
            in_code_block = False
            code_buffer = []
            code_lang = ""
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        if line.strip().startswith("> [!IMPORTANT]"):
            i += 1
            important_text = ""
            while i < len(lines) and lines[i].strip().startswith(">"):
                txt = lines[i].strip().lstrip("> ").strip()
                if txt: important_text += " " + txt
                i += 1
            if important_text.strip():
                add_important_box(doc, important_text.strip())
            continue

        if "|" in line and line.strip().startswith("|"):
            stripped = line.strip()
            if re.match(r"^\|[\s:|-]+\|$", stripped):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_header = stripped
                table_rows = []
            else:
                table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            add_table_from_md(doc, table_header, table_rows)
            in_table = False
            table_rows = []

        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith("## Table of Contents"):
            doc.add_heading(line[3:].strip(), level=1)
            p = doc.add_paragraph()
            # Adding TOC field to paragraph
            create_toc_field(p)
            doc.add_paragraph(" ")
            
            # Skip the manual TOC in the markdown
            i += 1
            while i < len(lines):
                line_text = lines[i].strip()
                if not line_text or re.match(r"^\d+\.", line_text):
                    i += 1
                else:
                    break
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        if line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/></w:pBdr>')
            pPr.append(pBdr)
            i += 1
            continue

        if line.startswith("- ") or line.startswith("* "):
            add_formatted_paragraph(doc, line[2:].strip(), style="List Bullet")
            i += 1
            continue

        num_match = re.match(r"^(\d+)\.\s+", line)
        if num_match:
            add_formatted_paragraph(doc, line[num_match.end():].strip(), style="List Number")
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        clean_line = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line)
        add_formatted_paragraph(doc, clean_line)
        i += 1

    if in_table:
        add_table_from_md(doc, table_header, table_rows)

    doc.save(OUT_PATH)
    print(f"[OK] Saved: {OUT_PATH}")

if __name__ == "__main__":
    convert_md_to_docx()
